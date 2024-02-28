import asyncio
import datetime
import json
import logging
import os
import re
import subprocess
import time
import math

import docx
import flet as ft
import psutil
import psycopg2 as pg
import qrcode
import requests
import xlrd
from docx2pdf import convert as pdf_convert
from psycopg2 import extras
from transliterate import translit
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from docxtpl import DocxTemplate
#from pypdf import PdfMerger

script_path = os.path.abspath(__file__)
script_directory = os.path.dirname(script_path)
os.chdir(script_directory)
project_folder = os.getcwd()

config_file_name = "config.json"

if not os.path.exists(os.path.join(project_folder, 'logs')):
    os.mkdir(os.path.join(project_folder, 'logs'))

logging.basicConfig(level=logging.INFO,
                    filename=f"{project_folder}/logs/{datetime.datetime.now().strftime('%d-%m-%Y-%H-%M')}.log",
                    filemode="w",
                    format="%(asctime)s %(levelname)s %(message)s",
                    encoding='utf-8'
                    )


def main(page: ft.Page):
    # ------------------НАЧАЛЬНЫЕ НАСТРОЙКИ------------------
    with open(os.path.join(project_folder, config_file_name), encoding='utf-8', mode="r") as config_file:
        config_data = json.load(config_file)

    control_data = config_data['control']
    url_base = control_data['base_url'].format(control_data['address'], control_data['port'])

    hf = ft.HapticFeedback()
    page.overlay.append(hf)
    page.title = "Connect"
    page.theme = ft.Theme(font_family="Montserrat")
    page.fonts = {
        "Montserrat": "fonts/Montserrat-SemiBold.ttf",
    }
    page.theme_mode = ft.ThemeMode.DARK

    def load_exe(exe_name: str):
        # Поиск и запуск исполняемых файлов

        logging.info(f"Поиск процесса {exe_name}.exe")
        if is_running(exe_name):
            logging.info(f"Процесс {exe_name}.exe обнаружен")
            if exe_name == 'Arena': change_screens('control')
        else:
            logging.info(f"Процесс {exe_name}.exe не обнаружен")
            if exe_name == 'Arena':
                page.dialog = dialog_resolume_start_menu
                dialog_resolume_start_menu.open = True
                page.update()

    def is_running(process_name: str):
        # Проверка наличия процесса в диспетчере задач

        for a in psutil.process_iter(['pid', 'name']):
            if a.name().split('.')[0] == process_name:
                return True
        return False

    def make_request(request_method: str, url: str, headers, data=None, files=None, params=None):
        # Отправка HTTP запроса

        response = requests.request(
            method=request_method,
            url=url,
            headers=headers,
            data=data,
            files=files,
            params=params
        )

        return response

    def get_headers(content_type: str):
        # Создание headers для HTTP запроса

        headers = control_data['request_headers']
        headers['Content-Type'] = content_type

        return headers

    def get_layer_clips(index: int):
        # Получение списка элементов слоя по индексу слоя

        url = f"{url_base}/layers/{index + 1}"
        api_response = make_request('GET', url, get_headers("application/json"))
        data = json.loads(api_response.text)
        print("layer: ", data['name']['value'], "id:", data['id'])
        return data['clips']

    def get_layers():
        # Получение списка слоёв проекта Resolume

        logging.info(f'[{get_layers.__name__}] Запрос на получение/обновление данных')

        api_response = make_request('GET', url_base, get_headers("application/json"))
        data = json.loads(api_response.text)['layers']

        logging.info(f'[{get_layers.__name__}] Данные получены')
        return data

    def power_off(e: ft.ControlEvent):
        dest = e.control.data
        if dest == 'layers_off':
            layers = get_layers()
            for index in range(len(layers)):
                url = f"{url_base}/layers/{index + 1}/clear"
                api_response = make_request('POST', url, headers=get_headers("application/json"))
                if api_response.status_code != 204:
                    close_dialog('e')
                    open_classic_snackbar("Ошибка при скрытии слоя", ft.colors.RED_ACCENT_200)
                    logging.error(f"[{power_off.__name__}] Ошибка при скрытии слоя")
                    logging.error(f"[{power_off.__name__}] URL: {url}")
                    send_error_message(
                        location=power_off.__name__,
                        error_text="Ошибка при скрытии слоя",
                        extra=url
                    )
                    break
                edit_control_card(index + 1, "---")

            close_dlg_power_menu()
            logging.info(f"[{power_off.__name__}] Все слои скрыты")
            # open_classic_snackbar("Все слои скрыты", ft.colors.GREEN)

    def open_clip(layer_index: int, clip_index: int, file: str):
        # Запуск нового клипа по индексу слоя и индексу клипа

        url = f"{url_base}/layers/{layer_index}/clips/{clip_index}/connect"
        logging.info(f'[{open_clip.__name__}] Запуск нового элемента {url}')
        api_response = make_request('POST', url, get_headers("application/json"))
        close_dialog('e')
        if api_response.status_code == 204:
            logging.info(f'[{open_clip.__name__}] Новый элемент открыт: {file}')
            open_classic_snackbar(f"Элемент {make_text_smaller(file)} загружен", ft.colors.GREEN_ACCENT_200)
            edit_control_card(layer_index, make_text_smaller(file), ft.colors.GREEN)
        else:
            logging.error(f'[{open_clip.__name__}] Ошибка при подключении к элементу: {api_response.text}')
            open_classic_snackbar("Ошибка при подключении к элементу", ft.colors.RED_ACCENT_200)
            send_error_message(
                location=open_clip.__name__,
                error_text=f"Ошибка при подключении к элементу: {api_response.text}",
                extra=url,
            )

    def find_connected(clips: {}):
        # Поиск активного элемента среди всех элементов слоя
        # [clip, clip_index] / False

        fl = False
        for clip_index in range(len(clips)):
            # print(clips[clip_index]['connected'])
            if clips[clip_index]['connected']['value'] == "Connected":
                fl = True
                return [clips[clip_index], clip_index]
        if not fl:
            return False

    def check_pin(e: ft.ControlEvent):
        # Проверка корректности пин-кода от Control

        user_pin = e.control.value
        if config_data['control']['pin'] == int(user_pin):
            change_screens('control')
        else:
            pin_field.border_color = ft.colors.ERROR
        page.update()

    def change_screens(target):
        # Переключение экранов

        page.clean()
        page.appbar = None
        page.navigation_bar = None
        page.scroll = None
        page.floating_action_button = None
        main_appbar.actions.clear()

        if target == "main":
            page.add(main_screen)
            page.appbar = main_appbar
            page.scroll = ft.ScrollMode.ADAPTIVE
            page.navigation_bar = navigation_bar
            page.vertical_alignment = ft.MainAxisAlignment.START,
            page.horizontal_alignment = ft.CrossAxisAlignment.CENTER
            main_appbar.leading = ft.IconButton(icon=ft.icons.LOCK_ROUNDED, on_click=lambda _: change_screens('login'))
        elif target == "login":
            page.appbar = None
            page.add(ft.Container(content=login_screen, expand=True))
            page.horizontal_alignment = ft.CrossAxisAlignment.CENTER
            open_classic_snackbar(f"Вы вышли из аккаунта", ft.colors.GREEN_ACCENT_200)
            main_appbar.title.value = "Connect"
            main_appbar.leading = None
            page.floating_action_button = ft.FloatingActionButton(
                icon=ft.icons.SETTINGS_REMOTE_ROUNDED, on_click=show_pin, bgcolor=ft.colors.BLUE_ACCENT_100
            )
        elif target == "edit_child":
            page.scroll = ft.ScrollMode.ADAPTIVE
            page.add(col_edit_child)
            page.appbar = edit_appbar
        elif target == "edit_mentor":
            page.scroll = ft.ScrollMode.ADAPTIVE
            page.add(col_edit_mentor)
            page.appbar = edit_appbar
        elif target == "control":
            logging.info('Выполнен переход к Control')
            dialog_enter_pin.open = False
            page.update()
            page.appbar = main_appbar
            main_appbar.title.value = "Control"
            main_appbar.actions = [ft.IconButton(ft.icons.SCREEN_SHARE_ROUNDED, on_click=open_proj_menu),
                                   ft.IconButton(ft.icons.RESTART_ALT_ROUNDED, on_click=lambda _: update_control()),
                                   ft.IconButton(ft.icons.POWER_SETTINGS_NEW_ROUNDED,
                                                 on_click=lambda _: open_power_menu_dialog()),
                                   ]
            main_appbar.leading = ft.IconButton(icon=ft.icons.ARROW_BACK_ROUNDED,
                                                on_click=lambda _: change_screens('login'))
            page.add(control_screen)
            update_control()
            page.scroll = ft.ScrollMode.ADAPTIVE

        page.update()

    pin_field = ft.TextField(text_size=30, password=True, keyboard_type=ft.KeyboardType.NUMBER,
                             text_align=ft.TextAlign.CENTER, on_submit=check_pin, border_width=2)

    def show_pin(e):
        hf.medium_impact()
        open_loading_snackbar("Загружаем данные")
        load_exe('Arena')

    def upload_new_element(layer_index, file):
        # Загрузка нового элемента

        logging.info(f"[{upload_new_element.__name__}] Загрузка нового элемента: {file}")
        layer_index += 1
        clips = get_layer_clips(layer_index - 1)

        for clip in range(len(clips)):
            if clips[clip]['connected']['value'] == "Empty":
                clip_index = clip + 1
                t = project_folder.replace("\\", "/")
                file_path = f"file:///{t}/assets/uploads/{file}"
                load_url = f"{url_base}/layers/{layer_index}/clips/{clip_index}/open"
                api_response = make_request('POST', load_url, get_headers('text/plain'),
                                            f"file:///{t}/assets/uploads/{file}")

                if api_response.status_code == 204:
                    logging.info(f"[{upload_new_element.__name__}] Элемент загружен в проект {file}")
                    open_clip(layer_index, clip_index, file)
                else:
                    close_dialog('e')
                    logging.error(f"[{upload_new_element.__name__}] Элемент не загружен в проект: {file}")
                    open_classic_snackbar("Элемент не загружен", ft.colors.RED_ACCENT_200)
                    send_error_message(
                        location=upload_new_element.__name__,
                        error_text=f"Элемент {file} не загружен в проект: {api_response.text}",
                        extra=f"{load_url}\n{file_path}",
                    )

                break

    def save_element(e: ft.ControlEvent):
        # Загрузка нового элемента в директорию

        logging.info(f"[{save_element.__name__}] Загрузка нового элемента в директорию")
        layer_index = e.control.data
        upload_list = []
        if element_picker.result is not None and element_picker.result.files is not None:
            open_loading_dialog()
            # open_loading_snackbar(f"Загружаем {make_text_smaller(element_picker.result.files[-1].name)}")
            for f in element_picker.result.files:
                upload_list.append(
                    ft.FilePickerUploadFile(
                        f.name,
                        upload_url=page.get_upload_url(f.name, 60),
                    )
                )
            element_picker.upload(upload_list)
            logging.info(f"[{save_element.__name__}] Загрузка нового элемента в директорию завершена")
            time.sleep(2)
            upload_new_element(layer_index, element_picker.result.files[-1].name)
        else:
            logging.info(f"[{save_element.__name__}] Загрузка нового элемента в директорию отменена")

    def on_uploading_element(e: ft.FilePickerUploadEvent):
        logging.info(f"[{save_element.__name__}] Загружено: {int(e.progress * 100)}%")
        dialog_loading.content.controls[2] = ft.Text(f"{int(e.progress * 100)}% / 100%", size=18)
        page.update()

    element_picker = ft.FilePicker(on_result=save_element, on_upload=on_uploading_element)

    def open_element_picker(e):
        hf.medium_impact()
        layer_index = e.control.data
        page.overlay.append(element_picker)
        page.update()
        element_picker.data = layer_index
        element_picker.pick_files(
            allow_multiple=False,
            file_type=ft.FilePickerFileType.CUSTOM,
            allowed_extensions=['jpg', 'jpeg', 'png', 'mp4', 'avi', 'mov']
        )

    def start_resolume():
        hf.medium_impact()
        time_to_open = config_data['control']['resolume_start_time']
        percents = 100 / time_to_open

        open_loading_dialog('Resolume Arena запускается')
        logging.info(f"Запуск процесса Arena.exe")
        subprocess.Popen(rf'{config_data["paths"]["arena"]}', creationflags=subprocess.CREATE_NEW_CONSOLE, shell=True)
        for i in range(1, time_to_open + 1):
            dialog_loading.content.controls[2] = ft.Text(f"{int(percents * i)}%", size=18)
            page.update()
            time.sleep(1)
        close_dialog('e')
        change_screens('control')

    def show_qr(e: ft.ControlEvent):
        phrase = e.control.data
        dialog_qr = ft.AlertDialog(
            content=ft.Image(src=f"assets/qrc/{phrase}.png", border_radius=ft.border_radius.all(10)),
            title=ft.Text("QR-код для регистрации", size=18),
            actions=[ft.Text(f"ключ-фраза: {phrase.split('_')[1]}", size=16)],
            actions_alignment=ft.MainAxisAlignment.CENTER
        )
        if phrase.split('_')[1] == "None":
            dialog_qr.content = ft.Text("Для данного пользователя не задана ключ-фраза", size=18)
            dialog_qr.actions.clear()
        else:
            qr_img = qrcode.make(data=f"https://t.me/crod_connect_bot?start={phrase}")
            qr_img.save(f"assets/qrc/{phrase}.png")

        page.dialog = dialog_qr
        dialog_qr.open = True
        page.update()

    def open_child(e: ft.ControlEvent):
        # фамилия, имя, рождения, коммент, группа
        data = e.control.data
        child_lastname.value = data['lastname']
        child_firstname.value = data['firstname']
        child_birth.value = data['birth']
        child_comment.value = data['comment']
        child_group_dropdown.value = data['group_num']
        child_parent.value = data['parent']
        child_parent_phone.value = data['parent_phone']

        col_edit_child.controls.pop()
        button_save_child.data = data['pass_phrase']
        col_edit_child.controls.append(
            ft.ResponsiveRow(
                controls=[
                    ft.FilledButton(
                        text="Показать QR-код", height=50, width=250,
                        style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder()),
                        col={"md": 4}, on_click=show_qr,
                        data=f"children_{data['pass_phrase']}"
                    ),
                    ft.ElevatedButton(
                        bgcolor=ft.colors.RED_ACCENT_200,
                        color=ft.colors.WHITE,
                        text="Исключить со смены", height=50, width=250,
                        style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder()),
                        col={"md": 4},
                        # on_click=remove_child
                    ),
                    button_save_child
                ]
            )
        )
        change_screens('edit_child')

    def open_mentor(e: ft.ControlEvent):
        data = e.control.data
        mentor_lastname.value = data['lastname']
        mentor_firstname.value = data['firstname']

        if len(col_edit_mentor.controls) > 3:
            col_edit_mentor.controls.pop()
        col_edit_mentor.controls.append(
            ft.ResponsiveRow(
                controls=[
                    ft.FilledButton(
                        text="Какая-то кнопка", height=50, width=250,
                        style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder()),
                        col={"md": 4}
                    ),
                    ft.FilledButton(
                        text="Показать QR-код", height=50, width=250,
                        style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder()),
                        col={"md": 4}, on_click=show_qr,
                        data=f"admins_{data['pass_phrase']}"
                    )
                ]
            )
        )
        change_screens('edit_mentor')

    def insert_metrics(filename, group_num):
        doc = DocxTemplate(filename)
        context = {"group_num": group_num,
                   "create_time": datetime.datetime.now().strftime("%d.%m.%Y в %H:%M"),
                   }
        doc.render(context)
        doc.save(filename)

    def generate_grouplist(e):
        open_loading_dialog("Генерируем документы")
        group_num = int(e.control.value)
        qr_filename_docx = f"{project_folder}/files/grouplists/Группа №{group_num} (QR-коды).docx"
        list_filename_docx = f"{project_folder}/files/grouplists/Группа №{group_num} (Список).docx"

        cur.execute(f"SELECT * FROM children WHERE group_num = '{group_num}'")
        data = cur.fetchall()

        arr_passphrases = []
        arr_child_info = []

        for child in data:
            qr_img = qrcode.make(f"https://t.me/crod_connect_bot?start=children_{child['pass_phrase']}")
            qr_img.save(f"assets/qrc/{child['pass_phrase']}.png")
            arr_passphrases.append(
                {'firstname': child['firstname'], 'lastname': child['lastname'], 'phrase': child['pass_phrase']})
            arr_child_info.append([
                child['lastname'],
                child['firstname'],
                datetime.datetime.strptime(child['birth'], '%Y-%m-%d').strftime('%d.%m.%Y'),
                child['comment'],
                " ".join([child['parent_lastname'], child['parent_firstname'], child['parent_middlename']]),
                child['parent_phone'],
            ])
        doc = docx.Document("assets/qr_template.docx")
        children_count = len(data)
        rows, cols = math.ceil(children_count / 4), 4
        table_qrc = doc.add_table(rows, cols)

        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

        index = 0
        for row in range(rows):
            for col in range(cols):
                if index + 1 <= len(arr_passphrases):
                    cell = table_qrc.cell(row, col)
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    paragraph.add_run().add_text(
                        f"{arr_passphrases[index]['firstname']} {arr_passphrases[index]['lastname']}")
                    paragraph.runs[-1].font.size = Pt(11)

                    image_path = f"assets/qrc/{arr_passphrases[index]['phrase']}.png"
                    paragraph.add_run().add_picture(image_path, width=Inches(1.7))
                    index += 1

        doc.save(qr_filename_docx)
        insert_metrics(qr_filename_docx, group_num)

        doc = docx.Document("assets/grouplist_template.docx")
        table = doc.tables[0]
        for i in range(children_count):
            row = table.add_row()

            for j, value in enumerate(arr_child_info[i]):
                cell = row.cells[j]
                cell.text = str(value)
                if j != 3:  # не выравнивается ячейка с особенностями
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.save(list_filename_docx)
        cur.execute(f"SELECT tid, pass_phrase FROM admins WHERE login = '{login_field.value}'")
        data = cur.fetchall()[0]['tid']
        insert_metrics(list_filename_docx, group_num)

        send_document([list_filename_docx, qr_filename_docx], data)
        close_dialog('e')
        open_classic_snackbar("Файлы отправлены в Telegram", ft.colors.GREEN)

    def open_group_picker():
        # Открытие диалога выбора группы для создания списков

        dialog_group_picker.content.controls[1].content = ft.Column(height=450)
        for a in range(1, config_data['group_count'] + 1):
            dialog_group_picker.content.controls[1].content.controls.append(ft.Radio(value=str(a), label=f"Группа № {a}"))
        page.dialog = dialog_group_picker
        dialog_group_picker.open = True
        page.update()

    def update_grouplist(e: ft.ControlEvent):
        open_loading_dialog()
        try:
            cur.execute(
                f"SELECT * FROM children WHERE group_num = '{e.control.value}'")
        except AttributeError:
            cur.execute(
                f"SELECT * FROM children WHERE group_num = '{e}'")
        data = cur.fetchall()
        col_groups.controls.clear()
        col_groups.controls.append(group_dropdown)
        for a in range(len(data)):
            col_groups.controls.append(
                ft.TextButton(
                    content=ft.Container(ft.Column(
                        [
                            ft.Text(
                                value=f"{a + 1}. {data[a]['lastname']} {data[a]['firstname']}",
                                size=18
                            )
                        ]
                    )),
                    on_click=open_child,
                    data={"lastname": data[a]['lastname'],
                          "firstname": data[a]['firstname'],
                          "birth": data[a]['birth'],
                          "comment": data[a]['comment'],
                          "group_num": data[a]['group_num'],
                          "pass_phrase": data[a]['pass_phrase'],
                          "parent": " ".join([data[a]['parent_lastname'], data[a]['parent_firstname'], data[a]['parent_middlename']]),
                          "parent_phone": f"+{data[a]['parent_phone']}"
                          }
                )
            )
            # col_groups.controls.append(ft.Divider(thickness=1))
        time.sleep(1)
        close_dialog('e')
        if len(data) == 0:
            open_classic_snackbar(f"Список группы №{e.control.value} отсутствует")
        page.update()

    def validate_login():
        if all([login_field.value, password_field.value]):
            button_login.disabled = False
        else:
            button_login.disabled = True
        page.update()

    def validate_child_info():
        if all([child_firstname.value, child_lastname.value]) and re.compile(r'^\d{4}-\d{1,2}-\d{1,2}$').match(
                child_birth.value):
            button_save_child.disabled = False
        else:
            button_save_child.disabled = True
        page.update()

    def login():
        login_field.value = login_field.value.strip()
        cur.execute(
            f"SELECT * FROM admins WHERE login = '{login_field.value}' AND password = '{password_field.value}' AND status = 'admin'")
        data = cur.fetchall()
        if len(data) != 0:

            change_screens('main')
            open_classic_snackbar(f"С возвращением, {data[0]['firstname']}", ft.colors.GREEN)
            update_grouplist(1)
            password_field.value = ""
            button_login.disabled = True
        else:
            open_classic_snackbar(f"Неверный логин или пароль", ft.colors.RED_ACCENT_200)
        page.update()

    def update_data():
        open_loading_dialog()
        files = [file for file in os.listdir("assets/uploads") if file.lower().endswith(('.xls', '.xlsx'))]
        cur.execute(f"SELECT lastname, firstname FROM admins WHERE login = '{login_field.value}'")
        data = cur.fetchall()[0]

        for file in files:
            name = file.split("_")[-2].split('.')[0]
            response = insert_data(f"assets/uploads/{file}", name)
            if response:
                config_data['tables_updates'][name]['date'] = datetime.datetime.now().strftime("%d.%m.%Y в %H:%M")
                config_data['tables_updates'][name]['user'] = f"{data['lastname']} {data['firstname'][0]}."
                with open(os.path.join(project_folder, config_file_name), encoding='utf-8', mode="w") as config_file:
                    json.dump(config_data, config_file, ensure_ascii=True, indent=2)
                time.sleep(2)
                update_tables_dates()
                open_classic_snackbar("Данные успешно загружены!", ft.colors.GREEN)
                os.remove(f"assets/uploads/{file}")
            else:
                open_classic_snackbar(f"Произошла ошибка при обновлении таблицы {name}", ft.colors.RED_ACCENT_200)
        close_dialog('e')

    def upload_tables(e):
        upload_list = []
        if tables_picker.result is not None and tables_picker.result.files is not None:
            for f in tables_picker.result.files:
                if True:
                    # if f.name.split(".")[0] in ['modules_info', 'mentors_info', 'children_info']:
                    upload_list.append(
                        ft.FilePickerUploadFile(
                            f.name,
                            upload_url=page.get_upload_url(f.name, 600),
                        )
                    )
                else:
                    open_classic_snackbar(f"Неизвестная таблица {f.name} не загружена", ft.colors.RED_ACCENT_200)
            open_loading_dialog()
            tables_picker.upload(upload_list)
            time.sleep(2)
            close_dialog('e')
            update_data()

    tables_picker = ft.FilePicker(on_result=upload_tables)

    def open_table_picker():
        page.overlay.append(tables_picker)
        page.update()
        tables_picker.pick_files(
            allow_multiple=True,
            file_type=ft.FilePickerFileType.CUSTOM,
            allowed_extensions=['xls', 'xlsx']
        )

    def update_child_info(e: ft.ControlEvent):
        open_loading_dialog()
        fl = False
        try:
            cur.execute(
                f"UPDATE children SET lastname = '{child_lastname.value}', firstname = '{child_firstname.value}', "
                f"birth = '{child_birth.value}' WHERE pass_phrase = '{e.control.data}'")
            connection.commit()
            fl = True
        except Exception as e:
            open_classic_snackbar(f"Ошибка при изменении данных: {e}", ft.colors.RED_ACCENT_200)
        if fl:
            open_classic_snackbar(f"Данные успешно обновлены!", ft.colors.GREEN)
        button_save_child.disabled = True
        close_dialog('e')
        page.go('/main')

    def update_welcome_text(e: ft.ControlEvent):
        config_data['welcome_text'] = e.control.value
        with open(os.path.join(project_folder, config_file_name), encoding='utf-8', mode="w") as config_file:
            json.dump(config_data, config_file, indent=2)

    def update_tables_dates():
        status_children.value = config_data['tables_updates']['children']['date']
        user_children.value = config_data['tables_updates']['children']['user']

        status_mentors.value = config_data['tables_updates']['mentors']['date']
        user_mentors.value = config_data['tables_updates']['mentors']['user']

        status_modules.value = config_data['tables_updates']['modules']['date']
        user_modules.value = config_data['tables_updates']['modules']['user']

        page.update()

    def send_table_example(login):
        cur.execute(f"SELECT tid, pass_phrase FROM admins WHERE login = '{login}'")
        data = cur.fetchall()[0]
        if data['tid'] is None:
            page.set_clipboard(f"https://t.me/crod_connect_bot?start=admins_{data['pass_phrase']}")
            open_classic_snackbar("Вы не зарегистрировались в боте ЦРОД.Connect, "
                                  "перейдите по ссылке, которая скопирована в буфер обмена, чтобы зарегистрироваться",
                                  ft.colors.RED_ACCENT_200)

        else:
            file_paths = [
                fr"{project_folder}\files\templates\_children_info.xlsx",
                fr"{project_folder}\files\templates\_modules_info.xlsx",
                fr"{project_folder}\files\templates\_mentors_info.xlsx"
            ]
            send_document(file_paths, data['tid'])

    def accept_change_group(e: ft.ControlEvent):
        old_group = e.control.data[0]
        new_group = child_group_dropdown.value
        phrase = e.control.data[1]
        page.dialog = dialog_change_group
        dialog_change_group.open = True
        page.update()

    def send_document(file_paths, tID):
        url = f'https://api.telegram.org/bot{config_data["bot"]["token"]}/sendDocument'
        for file_path in file_paths:
            print(file_path)
            with open(file_path, "rb") as file:
                files = {'document': (file_path, file)}
                data = {'chat_id': tID}
                response = requests.post(url=url, data=data, files=files)
                if response.status_code != 200:
                    open_classic_snackbar(f"Произошла ошибка при отправке документов: {response.text}",
                                          ft.colors.RED_ACCENT_200)

    def send_error_message(location: str, error_text: str, extra: str = "---"):
        url = f'https://api.telegram.org/bot{config_data["bot"]["token"]}/sendMessage'
        text = f"Ошибка\n\nМесто: {location}\nСодержание: {error_text}\n\nДополнительные данные: {extra}" \
            .replace("_", "\\_").replace("*", "\\*").replace("[", "\\[").replace("`", "\\`")
        data = {
            'chat_id': config_data['chat_ids']['errors'],
            'text': text,
            'parse_mode': 'Markdown'
        }
        api_response = make_request('POST', url, params=data, headers=None, files=None)
        if api_response.status_code != 200:
            logging.error(f"[{send_error_message.__name__}] Ошибка при отправке сообщения в Telegram")
            logging.error(f"[{send_error_message.__name__}] URL: {url}")
            logging.error(f"[{send_error_message.__name__}] Ошибка: {api_response.text}")

    def create_pass_phrase(lastname, firstname):
        ln = re.sub(r'[^a-zA-Zа-яА-ЯёЁ]', '', translit(lastname, language_code='ru', reversed=True))
        fn = re.sub(r'[^a-zA-Zа-яА-ЯёЁ]', '', translit(firstname, language_code='ru', reversed=True))
        date = datetime.datetime.now().strftime('%d%m%Y%H%M')
        return f"{ln}{fn}{date}"

    def insert_data(file_path, destination):
        wb = xlrd.open_workbook(file_path)
        ws = wb.sheet_by_index(0)
        rows_num = ws.nrows
        row = 1
        cur.execute(f"DELETE FROM admins WHERE status = '{destination}'")
        cur.execute(f"DELETE FROM children WHERE status = '{destination}'")
        while row < rows_num:
            # loading_text.value = f"Загрузка ({row} из {rows_num - 1})"
            page.update()
            if destination == "mentors":
                mentor = []
                for col in range(3):
                    mentor.append(ws.cell_value(row, col))
                # print(mentor)
                pass_phrase = create_pass_phrase(mentor[0], mentor[1])
                login = re.sub(r'[^a-zA-Zа-яА-ЯёЁ]', '', translit(mentor[0], language_code='ru', reversed=True))

                cur.execute(f"INSERT INTO admins (lastname, firstname, post, status, pass_phrase, login) "
                            f"VALUES ('{mentor[0]}', '{mentor[1]}', '{int(mentor[2])}', 'mentors', '{pass_phrase}', '{login}')")
            elif destination == "children":
                child = []
                for col in range(5):
                    child.append(ws.cell_value(row, col))
                # print(child)
                pass_phrase = create_pass_phrase(child[0], child[1])
                birth = xlrd.xldate.xldate_as_tuple(child[3], 0)
                birth = f"{birth[0]}-{birth[1]}-{birth[2]}"
                cur.execute(
                    f"INSERT INTO children (firstname, lastname, group_num, status, pass_phrase, birth, comment) "
                    f"VALUES ('{child[1]}', '{child[0]}', '{int(child[2])}', 'children', '{pass_phrase}', '{birth}', '{child[4]}')")

            connection.commit()
            row += 1
        return True

    def change_tab(e: ft.ControlEvent):
        hf.medium_impact()
        open_loading_dialog("Обновляем данные")
        try:
            cur_index = e.control.selected_index
            # main_appbar.title.value = config_data['tabs_names'][tab_indexes[cur_index]]
            main_appbar.actions.clear()
            page.update()
            update_tables_dates()
            col_groups.controls.clear()
            if cur_index == 0:  # группы
                update_grouplist(group_dropdown.value)
                group_dropdown.visible = True
                col_groups.visible = True

            elif cur_index == 1:  # модули
                open_loading_dialog()
                cur.execute(f"SELECT * FROM modules")
                close_dialog('e')

            elif cur_index == 2:  # воспиты
                cur.execute(
                    f"SELECT * FROM admins WHERE \"status\" = 'mentors'")
                data = cur.fetchall()
                if len(data) != 0:
                    for group in range(config_data['group_count']):
                        col_groups.controls.append(ft.Text(value=f"Группа №{group + 1}",
                                                           size=20,
                                                           text_align=ft.TextAlign.CENTER))
                        fl = False
                        for a in data:
                            if int(a['post']) == group + 1:
                                fl = True
                                col_groups.controls.append(
                                    ft.TextButton(
                                        content=ft.Container(ft.Column(
                                            [
                                                ft.Text(
                                                    value=f"{a['lastname']} {a['firstname']}",
                                                    size=18
                                                )
                                            ]
                                        )),
                                        on_click=open_mentor,
                                        style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder()),
                                        data={
                                            "lastname": a['lastname'],
                                            "firstname": a['firstname'],
                                            "pass_phrase": a['pass_phrase']
                                        },
                                    )
                                )
                        col_groups.controls.append(ft.Divider(thickness=1))
                        if not fl:
                            col_groups.controls.pop()
                            col_groups.controls.pop()
                    col_groups.visible = True
                else:
                    open_classic_snackbar(f"Список воспитателей отсутствует")
                group_dropdown.visible = False

            elif cur_index == 3:  # админка
                group_dropdown.visible = False
                cur.execute(f"SELECT * FROM admins WHERE \"status\" = 'admin'")
                data = cur.fetchall()
                col_groups.controls.clear()
                for a in data:
                    col_groups.controls.append(
                        ft.TextButton(
                            content=ft.Container(ft.Column(
                                [
                                    ft.Text(
                                        value=f"{a['lastname']} {a['firstname']}",
                                        size=18
                                    )
                                ]
                            )),
                            style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder())
                        )
                    )
                    col_groups.controls.append(ft.Divider(thickness=1))
                col_groups.visible = True
                if len(data) == 0:
                    open_classic_snackbar(f"Список администрации отсутствует")

            elif cur_index == 4:  # настройки
                main_appbar.actions = [ft.IconButton(icon=ft.icons.RESTART_ALT_ROUNDED)]
                col_groups.controls = [
                    ft.Container(ft.Text(value="Приветственное сообщение", size=18, ),
                                 margin=ft.margin.only(bottom=15)),
                    ft.TextField(text_align=ft.TextAlign.LEFT,
                                 multiline=True,
                                 value=config_data['welcome_text'].replace("*", ""),
                                 on_change=update_welcome_text
                                 ),
                    ft.Divider(thickness=1),
                    ft.Container(ft.Text(value="Данные о смене", size=18, ), margin=ft.margin.only(bottom=15)),
                    ft.Container(ft.Text(
                        value="Загрузите таблицы с внесёнными данными обо всех пользователях (можно загружать несколько таблиц сразу)."
                              "\nВнимание! Переименовывать таблицы нельзя!", size=14),
                        margin=ft.margin.only(bottom=15)),
                    ft.ResponsiveRow(
                        controls=[
                            ft.FilledButton("Выбрать файлы...", width=250,
                                            height=50,
                                            style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder()),
                                            col={"md": 4},
                                            on_click=lambda _: open_table_picker(),
                                            data="groups"
                                            ),
                            ft.ElevatedButton("Скачать шаблоны таблиц", width=250,
                                              height=50,
                                              style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder()),
                                              col={"md": 4},
                                              on_click=lambda _: send_table_example(login_field.value),
                                              ),
                            ft.ElevatedButton("Списки групп и QR", width=250,
                                              height=50,
                                              style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder()),
                                              col={"md": 4},
                                              on_click=lambda _: open_group_picker(),
                                              ),
                            ft.Divider(thickness=1),
                            ft.Container(ft.Text(value="Обновления таблиц", size=18, ),
                                         margin=ft.margin.only(bottom=15)),
                            files_status

                        ]
                    )
                ]
                col_groups.visible = True
            page.update()
            print("OK")
        except Exception as e:
            open_classic_snackbar(f"Ошибка в {change_tab.__name__}: {e}", ft.colors.RED_ACCENT_200)
        time.sleep(0.5)
        close_dialog('e')

    layers_info = {}

    rr = ft.ResponsiveRow()

    # ------------------ПАНЕЛИ CONTROL------------------
    main_panel = ft.ExpansionPanelList()

    screen_panel = ft.ExpansionPanel(
        expanded=False,
        header=ft.ListTile(title=ft.Text(f"HDMI")),
        content=ft.Column()
    )

    walls_panel = ft.ExpansionPanel(
        expanded=False,
        header=ft.ListTile(title=ft.Text(f"Стены")),
        content=ft.Column()
    )

    fons_panel = ft.ExpansionPanel(
        expanded=False,
        header=ft.ListTile(title=ft.Text(f"Фоны")),
        content=ft.Column()
    )

    logos_panel = ft.ExpansionPanel(
        expanded=False,
        header=ft.ListTile(title=ft.Text(f"Логотипы")),
        content=ft.Column()
    )

    layer_types = {
        # Соотношение панелей и названий из Resolume

        "WALLS": walls_panel,
        "FONS": fons_panel,
        "ЭКРАНЫ": screen_panel,
        "ЛОГОТИПЫ": logos_panel

    }

    layer_panel = {}

    def update_control():
        # Обновление экрана управления КЗ

        logging.info('Обновление данных Resolume')
        hf.medium_impact()
        open_loading_snackbar("Обновляем проект")
        rr.controls.clear()
        walls_panel.content.controls, fons_panel.content.controls, screen_panel.content.controls, logos_panel.content.controls = [], [], [], []
        main_panel.controls = [
            screen_panel,
            fons_panel,
            walls_panel,
            logos_panel
        ]
        control_screen.controls.clear()
        control_screen.controls.append(main_panel)

        layers_list = get_layers()

        for layer_index in range(len(layers_list) - 1, -1, -1):
            layer_name = layers_list[layer_index]['name']['value']
            current_panel = layer_types[layer_name[1:layer_name.find(']')]]
            layer_panel[layer_index + 1] = current_panel
            clips = get_layer_clips(layer_index)
            current_panel.content.controls.append(
                ft.Card(
                    ft.Column(
                        controls=[
                            ft.Container(
                                ft.Row(
                                    controls=[
                                        ft.Text(layers_list[layer_index]['name']['value'].split("]")[-1][1:], size=19,
                                                text_align=ft.TextAlign.CENTER)
                                    ],
                                    alignment=ft.MainAxisAlignment.CENTER,
                                ),
                                padding=20
                            ),
                            ft.Container(
                                content=ft.Row([ft.Text("---", size=16), ft.ProgressRing(visible=False, scale=0.5)],
                                               alignment=ft.MainAxisAlignment.CENTER, height=50),
                                padding=-15
                            ),
                            ft.Container(
                                ft.Row(
                                    controls=[
                                        ft.IconButton(ft.icons.ADD_ROUNDED, on_click=open_element_picker,
                                                      data=layer_index),
                                        ft.IconButton(icon=ft.icons.EDIT_ROUNDED, on_click=control_btn_pressed,
                                                      data=f"edit_{layer_index + 1}_{layers_list[layer_index]['id']}"
                                                      ),
                                        ft.IconButton(icon=ft.icons.KEYBOARD_ARROW_LEFT_OUTLINED,
                                                      on_click=control_btn_pressed,
                                                      data=f"prev_{layer_index + 1}_{layers_list[layer_index]['id']}"
                                                      ),
                                        ft.IconButton(icon=ft.icons.KEYBOARD_ARROW_RIGHT_OUTLINED,
                                                      on_click=control_btn_pressed,
                                                      data=f"next_{layer_index + 1}_{layers_list[layer_index]['id']}"
                                                      ),
                                        ft.IconButton(icon=ft.icons.VISIBILITY_OFF_OUTLINED,
                                                      on_click=control_btn_pressed,
                                                      data=f"stop_{layer_index + 1}_{layers_list[layer_index]['id']}"
                                                      ),
                                    ],
                                    alignment=ft.MainAxisAlignment.CENTER
                                ),
                                padding=10
                            )
                        ],
                        alignment=ft.MainAxisAlignment.START,
                        horizontal_alignment=ft.CrossAxisAlignment.CENTER
                    ),
                    data=layer_index

                )
            )

            connected_clip = find_connected(clips)
            if not connected_clip:
                layers_info[layers_list[layer_index]['id']] = -1
                current_panel.content.controls[-1].content.controls[1].content.controls[0].value = "---"
                current_panel.content.controls[-1].surface_tint_color = None
            else:
                layers_info[layers_list[layer_index]['id']] = connected_clip[1]

                current_panel.content.controls[-1].content.controls[1].content.controls[0].value = make_text_smaller(
                    connected_clip[0]['name']['value'])
                current_panel.content.controls[-1].surface_tint_color = ft.colors.GREEN
        time.sleep(2)
        page.update()
        open_classic_snackbar("Данные обновлены", ft.colors.GREEN_ACCENT_200)

    def control_btn_pressed(e: ft.ControlEvent):
        # Реакция на кнопки управления элементом слоя
        hf.medium_impact()
        request = e.control.data.split("_")
        layer_id = int(request[2])
        cur_clip_index = layers_info[layer_id]
        layer_index = int(request[1])
        action = request[0]

        print(request[0], layer_id, cur_clip_index, layer_index, action)
        if action in ['next', 'prev']:
            show_progress_ring(layer_index, True)
            if action == 'next':
                cur_clip_index += 1
            else:
                if cur_clip_index >= 1:
                    cur_clip_index -= 1
            layers_info[layer_id] = cur_clip_index
            url = f"{url_base}/layers/{layer_index}/clips/{cur_clip_index + 1}/connect"
            api_response = make_request('POST', url, get_headers("application/json"))
            if api_response.status_code != 204:
                logging.error(f"[{control_btn_pressed.__name__}] Ошибка при переключении элементов")
                logging.error(f"[{control_btn_pressed.__name__}] URL: {url}")
                open_classic_snackbar("Ошибка при переключении", ft.colors.RED_ACCENT_200)
                send_error_message(
                    location=control_btn_pressed.__name__,
                    error_text=f"Ошибка при переключении элементов",
                    extra=url,
                )
            else:
                time.sleep(2)
                clips = get_layer_clips(layer_index - 1)
                connected_clip = find_connected(clips)
                if not connected_clip:
                    pass
                    edit_control_card(layer_index, "Пустой слой")

                else:
                    connected_clip_info = connected_clip[0]
                    edit_control_card(layer_index, make_text_smaller(connected_clip_info['name']['value']),
                                      ft.colors.GREEN)
            show_progress_ring(layer_index, False)

        elif action == 'edit':
            clips = get_layer_clips(layer_index - 1)
            connected_clip = find_connected(clips)
            if not connected_clip:
                open_classic_snackbar("Сначала выберите элемент")
            else:
                dialog_edit.content = ft.Container(
                    content=ft.Column(
                        controls=[
                            ft.Row([ft.IconButton(ft.icons.ARROW_UPWARD_ROUNDED, scale=2,
                                                  on_click=edit_clip,
                                                  data=f"Position Y_minus_{connected_clip[0]['id']}"
                                                  )
                                    ],
                                   alignment=ft.MainAxisAlignment.CENTER),
                            ft.Row(
                                [ft.IconButton(ft.icons.ARROW_BACK_ROUNDED, scale=2,
                                               on_click=edit_clip,
                                               data=f"Position X_minus_{connected_clip[0]['id']}"
                                               ),
                                 ft.IconButton(ft.icons.FIT_SCREEN_ROUNDED, scale=2,
                                               on_click=edit_clip,
                                               data=f"center_center_{connected_clip[0]['id']}"
                                               ),
                                 ft.IconButton(ft.icons.ARROW_FORWARD_ROUNDED, scale=2,
                                               on_click=edit_clip,
                                               data=f"Position X_plus_{connected_clip[0]['id']}"
                                               )],
                                alignment=ft.MainAxisAlignment.CENTER,

                            ),
                            ft.Row([ft.IconButton(ft.icons.ARROW_DOWNWARD_ROUNDED, scale=2,
                                                  on_click=edit_clip,
                                                  data=f"Position Y_plus_{connected_clip[0]['id']}"
                                                  )],
                                   alignment=ft.MainAxisAlignment.CENTER),
                            ft.Divider(thickness=2),
                            ft.Row(
                                [
                                    ft.Text("Масштаб", width=100, size=16),
                                    ft.IconButton(ft.icons.ADD_ROUNDED,
                                                  on_click=edit_clip,
                                                  data=f"Scale_plus_{connected_clip[0]['id']}"
                                                  ),
                                    ft.IconButton(ft.icons.REMOVE_ROUNDED, on_click=edit_clip,
                                                  data=f"Scale_minus_{connected_clip[0]['id']}"
                                                  )
                                ],
                                alignment=ft.MainAxisAlignment.CENTER
                            ),
                            ft.Row(
                                [
                                    ft.Text("Ширина", width=100, size=16),
                                    ft.IconButton(ft.icons.ADD_ROUNDED,
                                                  on_click=edit_clip,
                                                  data=f"Scale W_plus_{connected_clip[0]['id']}"
                                                  ),
                                    ft.IconButton(ft.icons.REMOVE_ROUNDED, on_click=edit_clip,
                                                  data=f"Scale W_minus_{connected_clip[0]['id']}"
                                                  )
                                ],
                                alignment=ft.MainAxisAlignment.CENTER
                            ),
                            ft.Row(
                                [
                                    ft.Text("Высота", width=100, size=16),
                                    ft.IconButton(ft.icons.ADD_ROUNDED,
                                                  on_click=edit_clip,
                                                  data=f"Scale H_plus_{connected_clip[0]['id']}"
                                                  ),
                                    ft.IconButton(ft.icons.REMOVE_ROUNDED, on_click=edit_clip,
                                                  data=f"Scale H_minus_{connected_clip[0]['id']}"
                                                  )
                                ],
                                alignment=ft.MainAxisAlignment.CENTER
                            ),
                            # ft.Row([ft.Text("Наклон", width=100, size=16), ft.IconButton(ft.icons.ROTATE_LEFT_ROUNDED),
                            #         ft.IconButton(ft.icons.ROTATE_RIGHT_ROUNDED)], alignment=ft.MainAxisAlignment.CENTER)
                        ],
                        alignment=ft.MainAxisAlignment.CENTER,
                        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                        scroll=ft.ScrollMode.ADAPTIVE
                    ),
                    height=400,
                )
                page.dialog = dialog_edit
                dialog_edit.open = True
            page.update()

        elif action == 'stop':
            layers_info[layer_id] = -1
            url = f"{url_base}/layers/{layer_index}/clear"
            api_response = make_request('POST', url, get_headers("application/json"))
            if api_response.status_code == 204:
                edit_control_card(layer_index, "---")
            else:
                logging.error(f"[{control_btn_pressed.__name__}] Ошибка при отключении элемента")
                logging.error(f"[{control_btn_pressed.__name__}] URL: {url}")
                open_classic_snackbar("Ошибка при отключении", ft.colors.RED_ACCENT_200)
                send_error_message(
                    location=control_btn_pressed.__name__,
                    error_text="Ошибка при отключении элемента",
                    extra=url
                )
        page.update()

    def edit_control_card(layer_index: int, text: str, color=None):
        # Изменение свойств карточки слоя

        cur_panel = layer_panel[layer_index]
        cards = cur_panel.content.controls
        for i in range(len(cards)):
            if cards[i].data == layer_index - 1:
                cur_panel.content.controls[i].content.controls[1].content.controls[0].value = text
                cur_panel.content.controls[i].surface_tint_color = color
        page.update()

    def show_progress_ring(layer_index: int, action: bool):
        # Пока кольца загрузки вместо названия элемента
        # слоя во время переключения элементов

        cur_panel = layer_panel[layer_index]
        cards = cur_panel.content.controls
        for i in range(len(cards)):
            if cards[i].data == layer_index - 1:
                cur_panel.content.controls[i].content.controls[1].content.controls[1].visible = action
                cur_panel.content.controls[i].content.controls[1].content.controls[0].visible = not action
        page.update()

    def edit_clip(e: ft.ControlEvent):
        # Изменение свойств выбранного клипа

        hf.medium_impact()
        data = e.control.data.split("_")
        steps = {
            "Scale": 40,
            "Scale W": 10,
            "Scale H": 10,
            "Position X": 50,
            "Position Y": 50,
        }
        if data[0] != "center":
            step = steps[data[0]]
        clip_id = data[-1]
        url = f"{url_base}/clips/by-id/{clip_id}"
        old = json.loads(requests.get(url=url, headers=get_headers('application/json')).text)
        del old['video']['effects'][0]['bypassed']
        del old['video']['effects'][0]['mixer']
        del old['video']['sourceparams']
        if data[1] == "plus":
            old['video']['effects'][0]['params'][data[0]]['value'] += step
        elif data[1] == "minus":
            old['video']['effects'][0]['params'][data[0]]['value'] -= step
        elif data[1] == "center":
            old['video']['effects'][0]['params']['Position X']['value'] = 0
            old['video']['effects'][0]['params']['Position Y']['value'] = 0
        old = {"video": old['video']}
        requests.put(url=url, headers=get_headers('application/json'), data=json.dumps(old))

    def make_text_smaller(text):
        # Сокращение длины строки

        if len(text) > 15:
            text = text[:16] + "..."
        return text

    def open_proj_menu(e):
        # Меню проекторов

        hf.medium_impact()
        page.dialog = dialog_proj
        dialog_proj.open = True
        page.update()

    login_field = ft.TextField(label="Логин", text_align=ft.TextAlign.LEFT,
                               width=250, on_change=lambda _: validate_login(),
                               height=70)
    password_field = ft.TextField(label="Пароль", text_align=ft.TextAlign.LEFT,
                                  width=250, password=True, on_change=lambda _: validate_login(),
                                  can_reveal_password=True, height=70, on_submit=lambda _: login())
    button_login = ft.ElevatedButton("Войти", width=250, on_click=lambda _: login(), disabled=True, height=50,
                                     style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder()))

    button_save_child = ft.FilledButton(
        text="Сохранить", height=50, width=250,
        style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder(),
                             bgcolor=ft.colors.RED_ACCENT_200), col={"md": 4},
        disabled=True, on_click=update_child_info
    )

    group_dropdown = ft.Dropdown(
        visible=False,
        value="1",
        on_change=update_grouplist,
        label="Номер группы"
    )

    child_group_dropdown = ft.Dropdown(
        width=150,
        height=70,
        label="Группа",
        visible=False
    )

    for a in range(config_data['group_count']):
        group_dropdown.options.append(
            ft.dropdown.Option(str(a + 1))
        )
        child_group_dropdown.options.append(
            ft.dropdown.Option(str(a + 1))
        )

    main_appbar = ft.AppBar(title=ft.Text('Connect', size=18),
                            center_title=False,
                            bgcolor=ft.colors.SURFACE_VARIANT,
                            leading=None,
                            )

    edit_appbar = ft.AppBar(title=ft.Text('Изменение данных'),
                            center_title=True,
                            bgcolor=ft.colors.SURFACE_VARIANT,
                            leading=ft.IconButton(icon=ft.icons.ARROW_BACK_ROUNDED,
                                                  on_click=lambda _: change_screens('main'))
                            )

    col_groups = ft.Column(
        alignment=ft.MainAxisAlignment.START,
        horizontal_alignment=ft.CrossAxisAlignment.START,
        visible=False,
        scroll=ft.ScrollMode.ADAPTIVE,
        spacing=1

    )

    mentor_lastname = ft.TextField(label="Фамилия", height=70, width=250, col={"md": 4})
    mentor_firstname = ft.TextField(label="Имя", height=70, width=250, col={"md": 4})

    col_edit_mentor = ft.Column(
        alignment=ft.MainAxisAlignment.START,
        horizontal_alignment=ft.CrossAxisAlignment.START,
        scroll=ft.ScrollMode.ADAPTIVE,
        spacing=1,
        controls=[
            ft.Container(ft.Text(value="Персональная информация", size=18, ), margin=ft.margin.only(bottom=15)),
            ft.ResponsiveRow(
                controls=[
                    mentor_lastname,
                    mentor_firstname
                ]
            ),
            ft.Divider(thickness=1)
        ]
    )

    child_lastname = ft.TextField(label="Фамилия", height=70, width=250, col={"md": 4},
                                  on_change=lambda _: validate_child_info())
    child_firstname = ft.TextField(label="Имя", height=70, width=250, col={"md": 4},
                                   on_change=lambda _: validate_child_info())
    child_birth = ft.TextField(label="Дата рождения", height=70, width=250, col={"md": 4},
                               on_change=lambda _: validate_child_info())
    child_comment = ft.TextField(label="Особенности", width=350, col={"md": 4}, read_only=True,
                                 multiline=True)
    child_parent = ft.TextField(label="Имя", height=70, width=350, col={"md": 4},
                                read_only=True)
    child_parent_phone = ft.TextField(label="Телефон", height=70, width=250, col={"md": 4},
                                      read_only=True)



    col_edit_child = ft.Column(
        controls=[
            ft.Container(ft.Text(value="Данные о ребёнке", size=18, ), margin=ft.margin.only(bottom=15)),
            ft.ResponsiveRow(
                controls=[
                    child_lastname,
                    child_firstname,
                    child_birth,
                    child_comment,
                ],
            ),
            ft.Divider(thickness=1),
            ft.Container(ft.Text(value="Данные о родителе", size=18, ), margin=ft.margin.only(bottom=15)),
            ft.ResponsiveRow(
                controls=[
                    child_parent,
                    child_parent_phone,
                ],
            ),
            child_group_dropdown,
            ft.Divider(thickness=1),
            ft.Text(value="Действия", size=18),
        ],
        alignment=ft.MainAxisAlignment.START,
        horizontal_alignment=ft.CrossAxisAlignment.START,
        scroll=ft.ScrollMode.ADAPTIVE,
        spacing=1
    )

    status_children = ft.Text(config_data['tables_updates']['children']['date'])
    status_mentors = ft.Text(config_data['tables_updates']['mentors']['date'])
    status_modules = ft.Text(config_data['tables_updates']['modules']['date'])

    user_children = ft.Text(config_data['tables_updates']['children']['user'])
    user_mentors = ft.Text(config_data['tables_updates']['mentors']['user'])
    user_modules = ft.Text(config_data['tables_updates']['modules']['user'])

    files_status = ft.DataTable(
        columns=[
            ft.DataColumn(label=ft.Text("Таблица")),
            ft.DataColumn(label=ft.Text("Дата")),
            ft.DataColumn(label=ft.Text("Пользователь"))
        ],
        rows=[
            ft.DataRow(
                cells=[
                    ft.DataCell(ft.Text("Дети")),
                    ft.DataCell(status_children),
                    ft.DataCell(user_children)
                ]
            ),
            ft.DataRow(
                cells=[
                    ft.DataCell(ft.Text("Модули")),
                    ft.DataCell(status_modules),
                    ft.DataCell(user_modules)
                ]
            ),
            ft.DataRow(
                cells=[
                    ft.DataCell(ft.Text("Менторы")),
                    ft.DataCell(status_mentors),
                    ft.DataCell(user_mentors)
                ]
            ),

        ],
        column_spacing=10,
    )

    # ------------------ЭЛЕМЕНТЫ CONNECT------------------
    tab_indexes = {
        # Индексация экранов Connect

        0: "groups",
        1: "modules",
        2: "mentors",
        3: "admins",
        4: "settings",
    }

    navigation_bar = ft.NavigationBar(
        destinations=[
            ft.NavigationDestination(icon=ft.icons.GROUP_ROUNDED, label=config_data['tabs_names'][tab_indexes[0]]),
            ft.NavigationDestination(icon=ft.icons.BOOK_ROUNDED, label=config_data['tabs_names'][tab_indexes[1]]),
            ft.NavigationDestination(icon=ft.icons.SAFETY_DIVIDER_ROUNDED,
                                     label=config_data['tabs_names'][tab_indexes[2]]),
            ft.NavigationDestination(icon=ft.icons.HOME_WORK_ROUNDED, label=config_data['tabs_names'][tab_indexes[3]]),
            ft.NavigationDestination(icon=ft.icons.SETTINGS_ROUNDED, label=config_data['tabs_names'][tab_indexes[4]]),
        ],
        selected_index=None,
        on_change=change_tab,
        label_behavior=ft.NavigationBarLabelBehavior.ONLY_SHOW_SELECTED
    )

    # ------------------ДИАЛОГИ------------------
    dialog_group_picker = ft.AlertDialog(
        # Диалог выбора группы для создания списков
        title=ft.Text("Выбор группы"),
        content=ft.Column(
            height=200,
            controls=[
                ft.Text("Выберите, для какой группы необходимо создать списки", size=16),
                ft.RadioGroup(on_change=generate_grouplist)
            ]
        )
    )

    dialog_loading = ft.AlertDialog(
        # Диалог с кольцом загрузки
        title=ft.Text(""),
        content=ft.Column(
            controls=[
                ft.ProgressRing(),
                ft.Text("Загружаем", size=16),
                ft.Text("", size=14)
            ],
            alignment=ft.MainAxisAlignment.CENTER,
            horizontal_alignment=ft.CrossAxisAlignment.CENTER, height=100
        )
    )

    dialog_enter_pin = ft.AlertDialog(
        # Ввод кода от control_screen
        title=ft.Text("Введите ПИН-код"),
        actions=[
            pin_field
        ],
        actions_alignment=ft.MainAxisAlignment.CENTER,
        modal=True
    )

    dialog_power_menu = ft.AlertDialog(
        # Диалог с меню выключения
        title=ft.Text("Выключение"),
        content=ft.Column(
            controls=[
                ft.Text(
                    "Выберите требуемое действие",
                    size=16
                ),
            ],
            horizontal_alignment=ft.CrossAxisAlignment.START,
            scroll=ft.ScrollMode.ADAPTIVE
        ),
        actions=[
            ft.Column(
                controls=[
                    ft.ElevatedButton(
                        content=ft.Text(
                            "Показать слои",
                            size=16, color=ft.colors.WHITE
                        ),
                        width=350,
                        on_click=power_off,
                        data='layers_on'
                    ),
                    ft.ElevatedButton(
                        content=ft.Text(
                            "Скрыть слои",
                            size=16, color=ft.colors.WHITE
                        ),
                        width=350,
                        on_click=power_off,
                        data='layers_off'
                    ),
                    ft.Divider(thickness=2),
                    ft.ElevatedButton(
                        content=ft.Text(
                            "Выключить Resolume",
                            size=16,
                            color=ft.colors.WHITE),
                        width=350,
                        bgcolor=ft.colors.RED_500
                    ),
                    ft.ElevatedButton(
                        content=ft.Text(
                            "Выключить КЗ",
                            size=16,
                            color=ft.colors.WHITE
                        ),
                        width=350,
                        bgcolor=ft.colors.RED_500
                    ),
                ],
                horizontal_alignment=ft.CrossAxisAlignment.CENTER
            )
        ],
        actions_alignment=ft.MainAxisAlignment.CENTER
    )

    dialog_resolume_start_menu = ft.AlertDialog(
        # Диалог с кнопкой запуска Resolume
        modal=True,
        title=ft.Text("Resolume Arena"),
        content=ft.Text(
            "В данный момент Resolume Arena выключена. Включить?",
            size=16
        ),
        actions=[
            ft.ElevatedButton(
                text="Включить",
                on_click=lambda _: start_resolume(),
                color=ft.colors.GREEN_ACCENT_200
            ),
            ft.ElevatedButton(
                text="Назад",
                on_click=lambda _: close_dlg_resolume_start_menu(),
            )
        ],
        actions_alignment=ft.MainAxisAlignment.END
    )

    dialog_edit = ft.AlertDialog(
        # Диалог редактирования свойств элемента

        title=ft.Text("Свойства элемента"),
        actions=[
            ft.ElevatedButton(
                text="Закрыть",
                on_click=lambda _: close_dlg_edit()
            )
        ],
        actions_alignment=ft.MainAxisAlignment.END,

    )

    dialog_proj = ft.AlertDialog(
        # Диалог управления проекторами
        title=ft.Text("Проекторы"),
        actions=[
            ft.ElevatedButton(
                width=300,
                text="Запрос на включение",
                on_click=lambda _: close_dialog('e'),
                color=ft.colors.GREEN_ACCENT_200

            ),
            ft.ElevatedButton(
                width=300,
                text="Запрос на выключение",
                on_click=lambda _: close_dialog('e')
            )
        ],
        actions_alignment=ft.MainAxisAlignment.CENTER,
        content=ft.Text(
            "Если один/несколько проекторов отключились, нажмите на \"Запрос на включение\"",
            size=16
        )
    )

    dialog_change_group = ft.AlertDialog(
        title=ft.Text("Подтверждение перевода"),
        content=ft.Text("Подтвердите, что хотите перевести ребёнка в другую групу"),
        actions=[
            ft.TextButton("Подтвердить", on_click=lambda _: close_dialog('e')),
            ft.TextButton("Отклонить", on_click=lambda _: close_dialog('e'))
        ],
        modal=True
    )

    # ------------------ФУНКЦИИ ДИАЛОГОВ------------------

    def open_power_menu_dialog():
        hf.medium_impact()
        page.dialog = dialog_power_menu
        dialog_power_menu.open = True
        page.update()

    def open_loading_dialog(*args):
        if args:
            dialog_loading.content.controls[1].value = (args[-1])
        else:
            dialog_loading.content.controls[1].value = "Загрузка"
        page.dialog = dialog_loading
        dialog_loading.open = True
        page.update()

    def close_dlg_edit():
        hf.medium_impact()
        dialog_edit.open = False
        page.update()

    def close_dlg_resolume_start_menu():
        hf.medium_impact()
        dialog_resolume_start_menu.open = False
        page.update()

    def close_dlg_power_menu():
        dialog_power_menu.open = False
        page.update()

    def close_dialog(e: ft.ControlEvent):
        # Закрытие всех диалогов

        hf.medium_impact()
        dialog_loading.open = False
        dialog_enter_pin.open = False
        dialog_change_group.open = False
        dialog_proj.open = False
        time.sleep(2)
        page.update()

    # ------------------СНЕКБАРЫ------------------
    def open_loading_snackbar(text: str):
        # Бар с кольцом загрузки

        page.snack_bar = ft.SnackBar(
            content=ft.Row(
                controls=[
                    ft.ProgressRing(scale=0.5, color=ft.colors.BLACK),
                    ft.Text(text, size=16)
                ],
                alignment=ft.MainAxisAlignment.CENTER
            ),
            duration=1800
        )
        page.snack_bar.open = True
        page.update()

    def open_classic_snackbar(text: str, *args):
        # Классический бар

        page.snack_bar = ft.SnackBar(
            content=ft.Row(
                controls=[
                    ft.Text(text, size=16)
                ],
                alignment=ft.MainAxisAlignment.CENTER
            )
        )

        if args:
            page.snack_bar.bgcolor = args[0]

        page.snack_bar.open = True
        page.update()

    # ------------------ЭКРАНЫ------------------
    login_screen = ft.Column(
        # Экран авторизации

        controls=[
            ft.Container(ft.Image(src="/assets/logo2.png",
                                  fit=ft.ImageFit.CONTAIN,
                                  height=150,
                                  error_content=ft.ProgressRing()),
                         margin=ft.margin.only(bottom=20)),
            login_field,
            password_field,
            button_login
        ],
        alignment=ft.MainAxisAlignment.CENTER,
        horizontal_alignment=ft.CrossAxisAlignment.CENTER
    )

    main_screen = ft.Column(
        # Главный экран с панелями Connect

        controls=[
            col_groups
        ],
        alignment=ft.MainAxisAlignment.CENTER,
        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
    )

    control_screen = ft.Column(
        # Экран управления КЗ

        horizontal_alignment=ft.CrossAxisAlignment.CENTER,
        alignment=ft.MainAxisAlignment.START
    )

    # ------------------ТОЧКА ВХОДА------------------
    print(
        f"CROD Connect запущен: "
        f"http://localhost:{config_data['paths']['ngrok']['port']} "
        f"или https://{config_data['paths']['ngrok']['domain']}")
    change_screens('login')
    db_config = config_data['database']
    try:
        logging.info("Попытка подключения к БД")
        connection = pg.connect(
            host=db_config['host'],
            port=db_config['port'],
            user=db_config['user'],
            password=db_config['password'],
            database=db_config['db_name'],
        )
        cur = connection.cursor(cursor_factory=extras.RealDictCursor)
    except Exception as e:
        login_field.disabled, password_field.disabled = True, True
        # login_screen.controls.append(ft.Text("Устраните ошибку подключения к БД и перезагрузите страницу", size=16))
        page.update()
        open_classic_snackbar(f"Ошибка при подключении к БД", ft.colors.RED_ACCENT_200)
        logging.critical(f"Ошибка при подключении к БД : {e}")
        send_error_message(
            location="Подключение к БД",
            error_text=str(e)
        )


DEFAULT_FLET_PATH = ''

if __name__ == "__main__":

    with open(os.path.join(project_folder, config_file_name), encoding='utf-8', mode="r") as config_file:
        config_data = json.load(config_file)

    ngrok_open = False
    for a in psutil.process_iter(['pid', 'name']):
        if a.name().split('.')[0] == 'ngrok':
            ngrok_open = True
    if not ngrok_open:
        subprocess.Popen(
            fr'{config_data["paths"]["ngrok"]["path"]} http {config_data["paths"]["ngrok"]["port"]} --domain={config_data["paths"]["ngrok"]["domain"]}',
            creationflags=subprocess.CREATE_NEW_CONSOLE
        )
        logging.info("Процесс ngrok запущен")

    requests.get(url=f"https://{config_data['paths']['ngrok']['domain']}", headers={"ngrok-skip-browser-warning": "1"})

    flet_path = os.getenv("FLET_PATH", DEFAULT_FLET_PATH)
    ft.app(name=flet_path, target=main, view=None, port=config_data['paths']['ngrok']['port'],
           assets_dir="assets",
           upload_dir="assets/uploads")
